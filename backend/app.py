import io
import json
import os
import re
import threading
import time
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional
from uuid import uuid4

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from flask import Flask, Response, abort, jsonify, render_template, request, send_file, stream_with_context
from werkzeug.utils import secure_filename

try:
    from .history_store import HistoryStore
    from .document_vision_transcriber import DPI, postprocess_page_text, rasterize_page, transcribe_image
    from .translation_pipeline import load_translation_settings, translate_docx_file
except ImportError:
    from history_store import HistoryStore
    from document_vision_transcriber import DPI, postprocess_page_text, rasterize_page, transcribe_image
    from translation_pipeline import load_translation_settings, translate_docx_file

_RTL_RE = re.compile(r"[֐-׿؀-ۿݐ-ݿﭐ-﷿ﹰ-﻿]")

MODE_TRANSCRIBE = "transcribe"
MODE_TRANSCRIBE_TRANSLATE = "transcribe_translate"
MODE_TRANSLATE_ONLY = "translate_only"
SUPPORTED_MODES = {MODE_TRANSCRIBE, MODE_TRANSCRIBE_TRANSLATE, MODE_TRANSLATE_ONLY}


def _is_rtl(text: str) -> bool:
    return bool(_RTL_RE.search(text))


def _apply_paragraph_bidi(paragraph, is_rtl: bool) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1" if is_rtl else "0")
    p_pr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_rtl else WD_ALIGN_PARAGRAPH.LEFT


def _apply_run_bidi(run, is_rtl: bool) -> None:
    r_pr = run._r.get_or_add_rPr()
    rtl_element = OxmlElement("w:rtl")
    rtl_element.set(qn("w:val"), "1" if is_rtl else "0")
    r_pr.append(rtl_element)


BACKEND_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = BACKEND_DIR.parent
FRONTEND_DIR = PROJECT_ROOT / "frontend"
DATABASE_DIR = PROJECT_ROOT / "database"
UPLOAD_ROOT = DATABASE_DIR / "uploads"
JOB_OUTPUT_ROOT = DATABASE_DIR / "job_outputs"
HISTORY_PATH = DATABASE_DIR / "history.json"
TRANSLATION_SETTINGS_PATH = DATABASE_DIR / "translation_settings.json"
MAX_UPLOAD_SIZE = 100 * 1024 * 1024

app = Flask(
    __name__,
    template_folder=str(FRONTEND_DIR / "templates"),
    static_folder=str(FRONTEND_DIR / "static"),
)
app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_SIZE


@dataclass
class UploadRecord:
    upload_id: str
    original_name: str
    mode: str
    file_path: Path
    page_count: int
    preview_paths: Dict[int, Path]
    created_at: float = field(default_factory=time.time)


@dataclass
class JobRecord:
    job_id: str
    upload_id: str
    mode: str
    selected_pages: List[int]
    status: str = "queued"
    error: Optional[str] = None
    current_page: Optional[int] = None
    page_texts: Dict[int, str] = field(default_factory=dict)
    events: List[dict] = field(default_factory=list)
    artifacts: Dict[str, Path] = field(default_factory=dict)
    condition: threading.Condition = field(default_factory=threading.Condition)
    created_at: float = field(default_factory=time.time)
    started_at: Optional[float] = None
    completed_at: Optional[float] = None


uploads: Dict[str, UploadRecord] = {}
jobs: Dict[str, JobRecord] = {}
registry_lock = threading.Lock()
history_store = HistoryStore(HISTORY_PATH)


def ensure_storage() -> None:
    DATABASE_DIR.mkdir(parents=True, exist_ok=True)
    UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)
    JOB_OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)

    if not TRANSLATION_SETTINGS_PATH.exists():
        TRANSLATION_SETTINGS_PATH.write_text(
            json.dumps(
                {
                    "endpoint": "http://localhost:8020",
                    "model": "/data/models/gpt-oss-120b",
                    "api_key": "",
                    "batch_word_limit": 250,
                    "timeout": 180,
                    "max_retries": 5,
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )


def allowed_file_for_mode(filename: str, mode: str) -> bool:
    extension = Path(filename).suffix.lower()
    if mode in {MODE_TRANSCRIBE, MODE_TRANSCRIBE_TRANSLATE}:
        return extension == ".pdf"
    if mode == MODE_TRANSLATE_ONLY:
        return extension == ".docx"
    return False


def preview_url(upload_id: str, page_number: int) -> str:
    return f"/api/uploads/{upload_id}/preview/{page_number}"


def job_download_url(job_id: str, artifact_key: str) -> str:
    return f"/api/jobs/{job_id}/download/{artifact_key}"


def normalize_selected_pages(raw_pages: List[object], page_count: int) -> List[int]:
    selected = []
    seen = set()
    for value in raw_pages:
        try:
            page_number = int(value)
        except (TypeError, ValueError):
            continue
        if page_number < 1 or page_number > page_count or page_number in seen:
            continue
        selected.append(page_number)
        seen.add(page_number)
    return selected


def append_event(job: JobRecord, event_type: str, **payload: object) -> None:
    event = {
        "type": event_type,
        "timestamp": time.time(),
    }
    event.update(payload)
    with job.condition:
        job.events.append(event)
        job.condition.notify_all()


def get_upload(upload_id: str) -> UploadRecord:
    with registry_lock:
        record = uploads.get(upload_id)
    if record is None:
        abort(404, description="Upload not found.")
    return record


def get_job(job_id: str) -> JobRecord:
    with registry_lock:
        record = jobs.get(job_id)
    if record is None:
        abort(404, description="Job not found.")
    return record


def build_safe_stem(filename: str) -> str:
    stem = Path(filename).stem.strip()
    if not stem:
        return "document"
    return secure_filename(stem) or "document"


def build_download_map(job: JobRecord) -> Dict[str, str]:
    return {key: job_download_url(job.job_id, key) for key in sorted(job.artifacts)}


def get_job_snapshot(job: JobRecord) -> dict:
    with job.condition:
        snapshot = {
            "job_id": job.job_id,
            "upload_id": job.upload_id,
            "mode": job.mode,
            "status": job.status,
            "error": job.error,
            "current_page": job.current_page,
            "selected_pages": list(job.selected_pages),
            "completed_pages": sorted(job.page_texts.keys()),
            "download_urls": build_download_map(job) if job.status == "completed" else None,
        }
    return snapshot


def _write_txt_zip(path: Path, selected_pages: List[int], page_texts: Dict[int, str]) -> None:
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zip_file:
        for page_number in selected_pages:
            page_text = page_texts.get(page_number, "")
            zip_file.writestr(f"page_{page_number:03d}.txt", page_text)


def _build_transcription_docx_buffer(selected_pages: List[int], page_texts: Dict[int, str]) -> io.BytesIO:
    document = Document()
    for index, page_number in enumerate(selected_pages):
        lines = page_texts.get(page_number, "").split("\n")
        last_paragraph = None
        for line in lines:
            paragraph = document.add_paragraph()
            rtl = _is_rtl(line)
            _apply_paragraph_bidi(paragraph, rtl)
            if line:
                run = paragraph.add_run(line)
                _apply_run_bidi(run, rtl)
            last_paragraph = paragraph
        if index < len(selected_pages) - 1 and last_paragraph is not None:
            last_paragraph.add_run().add_break(WD_BREAK.PAGE)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _mark_job_completed(job: JobRecord) -> None:
    with job.condition:
        job.status = "completed"
        job.current_page = None
        job.completed_at = time.time()
    append_event(job, "status", status="completed", download_urls=build_download_map(job))


def _mark_job_failed(job: JobRecord, error_message: str) -> None:
    with job.condition:
        job.status = "failed"
        job.error = error_message
        job.current_page = None
        job.completed_at = time.time()
    append_event(job, "status", status="failed", error=error_message)


def _record_history(job: JobRecord, upload: UploadRecord) -> None:
    history_store.append(
        {
            "job_id": job.job_id,
            "upload_id": upload.upload_id,
            "filename": upload.original_name,
            "mode": job.mode,
            "status": job.status,
            "error": job.error,
            "selected_pages": list(job.selected_pages),
            "page_count": upload.page_count,
            "created_at": job.created_at,
            "started_at": job.started_at,
            "completed_at": job.completed_at,
            "download_urls": build_download_map(job) if job.status == "completed" else {},
        }
    )


def _run_translate_only_job(job: JobRecord, upload: UploadRecord, output_dir: Path) -> None:
    translated_path = output_dir / f"{build_safe_stem(upload.original_name)}_translated_en.docx"

    settings = load_translation_settings(TRANSLATION_SETTINGS_PATH)

    def on_progress(payload: dict) -> None:
        append_event(job, "translation_progress", **payload)

    append_event(job, "translation_start", filename=upload.original_name)
    report = translate_docx_file(upload.file_path, translated_path, settings, progress_callback=on_progress)
    job.artifacts["translated_docx"] = translated_path
    append_event(job, "translation_complete", **report)


def _run_transcription_modes(job: JobRecord, upload: UploadRecord, output_dir: Path) -> None:
    doc = None
    docx_path = output_dir / f"{build_safe_stem(upload.original_name)}_transcription.docx"
    txt_zip_path = output_dir / f"{build_safe_stem(upload.original_name)}_pages_txt.zip"

    try:
        doc = fitz.open(upload.file_path)
        total_pages = len(job.selected_pages)
        for selected_index, page_number in enumerate(job.selected_pages, start=1):
            with job.condition:
                job.current_page = page_number
            append_event(
                job,
                "page_start",
                page_number=page_number,
                selected_index=selected_index,
                total_selected=total_pages,
                preview_url=preview_url(upload.upload_id, page_number),
            )

            jpeg_bytes = rasterize_page(doc[page_number - 1], dpi=DPI)

            def on_token(token: str, page_ref: int = page_number) -> None:
                append_event(job, "token", page_number=page_ref, token=token)

            text = postprocess_page_text(
                transcribe_image(jpeg_bytes, on_token=on_token, print_tokens=False)
            )
            with job.condition:
                job.page_texts[page_number] = text
            append_event(
                job,
                "page_complete",
                page_number=page_number,
                preview_url=preview_url(upload.upload_id, page_number),
                text=text,
            )

        _write_txt_zip(txt_zip_path, job.selected_pages, job.page_texts)
        docx_path.write_bytes(_build_transcription_docx_buffer(job.selected_pages, job.page_texts).getvalue())
        job.artifacts["txt"] = txt_zip_path
        job.artifacts["docx"] = docx_path

        if job.mode == MODE_TRANSCRIBE_TRANSLATE:
            translated_path = output_dir / f"{build_safe_stem(upload.original_name)}_translated_en.docx"
            settings = load_translation_settings(TRANSLATION_SETTINGS_PATH)

            def on_progress(payload: dict) -> None:
                append_event(job, "translation_progress", **payload)

            append_event(job, "translation_start", filename=docx_path.name)
            report = translate_docx_file(docx_path, translated_path, settings, progress_callback=on_progress)
            job.artifacts["translated_docx"] = translated_path
            append_event(job, "translation_complete", **report)
    finally:
        if doc is not None:
            doc.close()


def run_processing_job(job_id: str) -> None:
    job = get_job(job_id)
    upload = get_upload(job.upload_id)
    output_dir = JOB_OUTPUT_ROOT / job.job_id
    output_dir.mkdir(parents=True, exist_ok=True)

    with job.condition:
        job.status = "running"
        job.started_at = time.time()
    append_event(job, "status", status="running", mode=job.mode)

    try:
        if job.mode == MODE_TRANSLATE_ONLY:
            _run_translate_only_job(job, upload, output_dir)
        else:
            _run_transcription_modes(job, upload, output_dir)
        _mark_job_completed(job)
    except Exception as exc:
        _mark_job_failed(job, str(exc))
    finally:
        _record_history(job, upload)


@app.get("/")
def index() -> str:
    return render_template("index.html")


@app.get("/history")
def history_page() -> str:
    return render_template("history.html")


@app.get("/favicon.ico")
def favicon() -> Response:
    return Response(status=204)


@app.get("/api/history")
def api_history() -> Response:
    return jsonify({"items": history_store.list()})


@app.post("/api/upload")
def upload_document() -> Response:
    mode = str(request.form.get("mode") or MODE_TRANSCRIBE).strip().lower()
    if mode not in SUPPORTED_MODES:
        return jsonify({"error": "Unsupported mode."}), 400

    uploaded_file = request.files.get("document") or request.files.get("pdf")
    if uploaded_file is None or not uploaded_file.filename:
        if mode == MODE_TRANSLATE_ONLY:
            return jsonify({"error": "Choose a DOCX file to upload."}), 400
        return jsonify({"error": "Choose a PDF file to upload."}), 400

    original_name = uploaded_file.filename
    if not allowed_file_for_mode(original_name, mode):
        if mode == MODE_TRANSLATE_ONLY:
            return jsonify({"error": "Translate-only mode supports DOCX uploads only."}), 400
        return jsonify({"error": "Selected mode supports PDF uploads only."}), 400

    upload_id = uuid4().hex
    upload_dir = UPLOAD_ROOT / upload_id
    preview_dir = upload_dir / "previews"
    upload_dir.mkdir(parents=True, exist_ok=True)
    preview_dir.mkdir(parents=True, exist_ok=True)

    safe_name = secure_filename(original_name) or (
        "upload.docx" if mode == MODE_TRANSLATE_ONLY else "upload.pdf"
    )
    extension = ".docx" if mode == MODE_TRANSLATE_ONLY else ".pdf"
    if Path(safe_name).suffix.lower() != extension:
        safe_name = f"{Path(safe_name).stem or 'upload'}{extension}"
    file_path = upload_dir / safe_name
    uploaded_file.save(file_path)

    page_count = 0
    preview_paths: Dict[int, Path] = {}
    pages: List[dict] = []

    if mode in {MODE_TRANSCRIBE, MODE_TRANSCRIBE_TRANSLATE}:
        doc = None
        try:
            doc = fitz.open(file_path)
            page_count = len(doc)
            if page_count == 0:
                return jsonify({"error": "The uploaded PDF has no pages."}), 400

            for page_index in range(page_count):
                page_number = page_index + 1
                preview_path = preview_dir / f"page_{page_number:03d}.jpg"
                preview_path.write_bytes(rasterize_page(doc[page_index], dpi=96))
                preview_paths[page_number] = preview_path
                pages.append(
                    {
                        "page_number": page_number,
                        "preview_url": preview_url(upload_id, page_number),
                    }
                )
        except Exception as exc:
            return jsonify({"error": f"Could not open the uploaded PDF: {exc}"}), 400
        finally:
            if doc is not None:
                doc.close()

    record = UploadRecord(
        upload_id=upload_id,
        original_name=original_name,
        mode=mode,
        file_path=file_path,
        page_count=page_count,
        preview_paths=preview_paths,
    )
    with registry_lock:
        uploads[upload_id] = record

    return jsonify(
        {
            "upload_id": upload_id,
            "filename": original_name,
            "mode": mode,
            "page_count": page_count,
            "pages": pages,
        }
    )


@app.get("/api/uploads/<upload_id>/preview/<int:page_number>")
def serve_preview(upload_id: str, page_number: int):
    upload = get_upload(upload_id)
    preview_path = upload.preview_paths.get(page_number)
    if preview_path is None or not preview_path.exists():
        abort(404, description="Preview not found.")
    return send_file(preview_path, mimetype="image/jpeg", max_age=0)


@app.post("/api/jobs")
def create_job() -> Response:
    payload = request.get_json(silent=True) or {}
    upload_id = payload.get("upload_id")
    raw_pages = payload.get("selected_pages") or []

    if not upload_id:
        return jsonify({"error": "Missing upload id."}), 400

    upload = get_upload(str(upload_id))

    selected_pages: List[int]
    if upload.mode in {MODE_TRANSCRIBE, MODE_TRANSCRIBE_TRANSLATE}:
        selected_pages = normalize_selected_pages(list(raw_pages), upload.page_count)
        if not selected_pages:
            return jsonify({"error": "Select at least one page."}), 400
    else:
        selected_pages = []

    job_id = uuid4().hex
    job = JobRecord(
        job_id=job_id,
        upload_id=upload.upload_id,
        mode=upload.mode,
        selected_pages=selected_pages,
    )
    with registry_lock:
        jobs[job_id] = job

    append_event(job, "status", status="queued", mode=job.mode)
    worker = threading.Thread(target=run_processing_job, args=(job_id,), daemon=True)
    worker.start()

    return jsonify(
        {
            "job_id": job_id,
            "stream_url": f"/api/jobs/{job_id}/stream",
            "status_url": f"/api/jobs/{job_id}",
            "mode": job.mode,
        }
    )


@app.get("/api/jobs/<job_id>")
def job_status(job_id: str) -> Response:
    return jsonify(get_job_snapshot(get_job(job_id)))


@app.get("/api/jobs/<job_id>/stream")
def stream_job(job_id: str):
    job = get_job(job_id)

    @stream_with_context
    def event_stream():
        index = 0
        while True:
            event = None
            finished = False
            with job.condition:
                while index >= len(job.events) and job.status not in {"completed", "failed"}:
                    job.condition.wait(timeout=15)
                    if index >= len(job.events) and job.status not in {"completed", "failed"}:
                        break
                if index < len(job.events):
                    event = job.events[index]
                    index += 1
                else:
                    finished = job.status in {"completed", "failed"}

            if event is not None:
                yield f"data: {json.dumps(event)}\n\n"
                continue

            if finished:
                break

            yield ": keep-alive\n\n"

    response = Response(event_stream(), mimetype="text/event-stream")
    response.headers["Cache-Control"] = "no-cache"
    response.headers["X-Accel-Buffering"] = "no"
    return response


@app.get("/api/jobs/<job_id>/download/<artifact_key>")
def download_artifact(job_id: str, artifact_key: str):
    job = get_job(job_id)
    if job.status != "completed":
        abort(409, description="The job is not finished yet.")

    artifact_path = job.artifacts.get(artifact_key)
    if artifact_path is None or not artifact_path.exists():
        abort(404, description="Requested artifact is not available.")

    upload = get_upload(job.upload_id)
    stem = build_safe_stem(upload.original_name)

    if artifact_key == "txt":
        return send_file(
            artifact_path,
            mimetype="application/zip",
            as_attachment=True,
            download_name=f"{stem}_pages_txt.zip",
        )
    if artifact_key == "docx":
        return send_file(
            artifact_path,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"{stem}_transcription.docx",
        )
    if artifact_key == "translated_docx":
        return send_file(
            artifact_path,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"{stem}_translated_en.docx",
        )

    abort(404, description="Unsupported artifact type.")


@app.get("/api/jobs/<job_id>/download/txt")
def download_txt(job_id: str):
    return download_artifact(job_id, "txt")


@app.get("/api/jobs/<job_id>/download/docx")
def download_docx(job_id: str):
    return download_artifact(job_id, "docx")


ensure_storage()


if __name__ == "__main__":
    debug_enabled = os.environ.get("FLASK_DEBUG", "0") == "1"
    port = int(os.environ.get("PORT", "5000"))
    app.run(
        host="127.0.0.1",
        port=port,
        debug=debug_enabled,
        use_reloader=debug_enabled,
        threaded=True,
    )