import io
import json
import os
import re
import threading
import time
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional
from uuid import uuid4

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from flask import Flask, Response, abort, jsonify, render_template, request, send_file, stream_with_context
from werkzeug.utils import secure_filename

_RTL_RE = re.compile(r"[֐-׿؀-ۿݐ-ݿﭐ-﷿ﹰ-﻿]")


def _is_rtl(text: str) -> bool:
    return bool(_RTL_RE.search(text))


def _apply_paragraph_bidi(paragraph, is_rtl: bool) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1" if is_rtl else "0")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_rtl else WD_ALIGN_PARAGRAPH.LEFT


def _apply_run_bidi(run, is_rtl: bool) -> None:
    rPr = run._r.get_or_add_rPr()
    rtl_el = OxmlElement("w:rtl")
    rtl_el.set(qn("w:val"), "1" if is_rtl else "0")
    rPr.append(rtl_el)

from ocr_with_ollama import DPI, postprocess_page_text, rasterize_page, transcribe_image


BASE_DIR = Path(__file__).resolve().parent
INSTANCE_DIR = BASE_DIR / "instance"
UPLOAD_ROOT = INSTANCE_DIR / "uploads"
MAX_UPLOAD_SIZE = 100 * 1024 * 1024

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_SIZE


@dataclass
class UploadRecord:
    upload_id: str
    original_name: str
    pdf_path: Path
    page_count: int
    preview_paths: Dict[int, Path]
    created_at: float = field(default_factory=time.time)


@dataclass
class JobRecord:
    job_id: str
    upload_id: str
    selected_pages: List[int]
    status: str = "queued"
    error: Optional[str] = None
    current_page: Optional[int] = None
    page_texts: Dict[int, str] = field(default_factory=dict)
    events: List[dict] = field(default_factory=list)
    condition: threading.Condition = field(default_factory=threading.Condition)
    created_at: float = field(default_factory=time.time)
    started_at: Optional[float] = None
    completed_at: Optional[float] = None


uploads: Dict[str, UploadRecord] = {}
jobs: Dict[str, JobRecord] = {}
registry_lock = threading.Lock()


def ensure_storage() -> None:
    UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)


def allowed_pdf(filename: str) -> bool:
    return Path(filename).suffix.lower() == ".pdf"


def preview_url(upload_id: str, page_number: int) -> str:
    return f"/api/uploads/{upload_id}/preview/{page_number}"


def job_download_url(job_id: str, fmt: str) -> str:
    return f"/api/jobs/{job_id}/download/{fmt}"


def normalize_selected_pages(raw_pages: List[object], page_count: int) -> List[int]:
    selected = []
    seen = set()
    for value in raw_pages:
        try:
            page_number = int(value)
        except (TypeError, ValueError):
            continue
        if page_number < 1 or page_number > page_count or page_number in seen:
            continue
        selected.append(page_number)
        seen.add(page_number)
    return selected


def append_event(job: JobRecord, event_type: str, **payload: object) -> None:
    event = {
        "type": event_type,
        "timestamp": time.time(),
    }
    event.update(payload)
    with job.condition:
        job.events.append(event)
        job.condition.notify_all()


def get_upload(upload_id: str) -> UploadRecord:
    with registry_lock:
        record = uploads.get(upload_id)
    if record is None:
        abort(404, description="Upload not found.")
    return record


def get_job(job_id: str) -> JobRecord:
    with registry_lock:
        record = jobs.get(job_id)
    if record is None:
        abort(404, description="Job not found.")
    return record


def get_job_snapshot(job: JobRecord) -> dict:
    with job.condition:
        snapshot = {
            "job_id": job.job_id,
            "upload_id": job.upload_id,
            "status": job.status,
            "error": job.error,
            "current_page": job.current_page,
            "selected_pages": list(job.selected_pages),
            "completed_pages": sorted(job.page_texts.keys()),
            "download_urls": None,
        }
    if snapshot["status"] == "completed":
        snapshot["download_urls"] = {
            "txt": job_download_url(job.job_id, "txt"),
            "docx": job_download_url(job.job_id, "docx"),
        }
    return snapshot


def build_safe_stem(filename: str) -> str:
    stem = Path(filename).stem.strip()
    if not stem:
        return "transcription"
    return secure_filename(stem) or "transcription"


def run_transcription_job(job_id: str) -> None:
    job = get_job(job_id)
    upload = get_upload(job.upload_id)

    with job.condition:
        job.status = "running"
        job.started_at = time.time()
    append_event(job, "status", status="running")

    doc = None
    try:
        doc = fitz.open(upload.pdf_path)
        total_pages = len(job.selected_pages)
        for selected_index, page_number in enumerate(job.selected_pages, start=1):
            with job.condition:
                job.current_page = page_number
            append_event(
                job,
                "page_start",
                page_number=page_number,
                selected_index=selected_index,
                total_selected=total_pages,
                preview_url=preview_url(upload.upload_id, page_number),
            )

            jpeg_bytes = rasterize_page(doc[page_number - 1], dpi=DPI)

            def on_token(token: str, page_ref: int = page_number) -> None:
                append_event(job, "token", page_number=page_ref, token=token)

            text = postprocess_page_text(
                transcribe_image(jpeg_bytes, on_token=on_token, print_tokens=False)
            )
            with job.condition:
                job.page_texts[page_number] = text
            append_event(job, "page_complete", page_number=page_number, text=text)

        with job.condition:
            job.status = "completed"
            job.current_page = None
            job.completed_at = time.time()
        append_event(
            job,
            "status",
            status="completed",
            download_urls={
                "txt": job_download_url(job.job_id, "txt"),
                "docx": job_download_url(job.job_id, "docx"),
            },
        )
    except Exception as exc:
        with job.condition:
            job.status = "failed"
            job.error = str(exc)
            job.current_page = None
            job.completed_at = time.time()
        append_event(job, "status", status="failed", error=str(exc))
    finally:
        if doc is not None:
            doc.close()


@app.get("/")
def index() -> str:
    return render_template("index.html")


@app.get("/favicon.ico")
def favicon() -> Response:
    return Response(status=204)


@app.post("/api/upload")
def upload_pdf() -> Response:
    uploaded_file = request.files.get("pdf")
    if uploaded_file is None or not uploaded_file.filename:
        return jsonify({"error": "Choose a PDF file to upload."}), 400

    original_name = uploaded_file.filename
    safe_name = secure_filename(original_name) or "upload.pdf"
    if not allowed_pdf(original_name):
        return jsonify({"error": "Only PDF uploads are supported."}), 400
    if Path(safe_name).suffix.lower() != ".pdf":
        safe_name = f"{Path(safe_name).stem or 'upload'}.pdf"

    upload_id = uuid4().hex
    upload_dir = UPLOAD_ROOT / upload_id
    preview_dir = upload_dir / "previews"
    upload_dir.mkdir(parents=True, exist_ok=True)
    preview_dir.mkdir(parents=True, exist_ok=True)

    pdf_path = upload_dir / safe_name
    uploaded_file.save(pdf_path)

    doc = None
    try:
        doc = fitz.open(pdf_path)
        page_count = len(doc)
        if page_count == 0:
            return jsonify({"error": "The uploaded PDF has no pages."}), 400

        preview_paths = {}
        pages = []
        for page_index in range(page_count):
            page_number = page_index + 1
            preview_path = preview_dir / f"page_{page_number:03d}.jpg"
            preview_path.write_bytes(rasterize_page(doc[page_index], dpi=96))
            preview_paths[page_number] = preview_path
            pages.append(
                {
                    "page_number": page_number,
                    "preview_url": preview_url(upload_id, page_number),
                }
            )
    except Exception as exc:
        return jsonify({"error": f"Could not open the uploaded PDF: {exc}"}), 400
    finally:
        if doc is not None:
            doc.close()

    record = UploadRecord(
        upload_id=upload_id,
        original_name=original_name,
        pdf_path=pdf_path,
        page_count=page_count,
        preview_paths=preview_paths,
    )
    with registry_lock:
        uploads[upload_id] = record

    return jsonify(
        {
            "upload_id": upload_id,
            "filename": original_name,
            "page_count": page_count,
            "pages": pages,
        }
    )


@app.get("/api/uploads/<upload_id>/preview/<int:page_number>")
def serve_preview(upload_id: str, page_number: int):
    upload = get_upload(upload_id)
    preview_path = upload.preview_paths.get(page_number)
    if preview_path is None or not preview_path.exists():
        abort(404, description="Preview not found.")
    return send_file(preview_path, mimetype="image/jpeg", max_age=0)


@app.post("/api/jobs")
def create_job() -> Response:
    payload = request.get_json(silent=True) or {}
    upload_id = payload.get("upload_id")
    raw_pages = payload.get("selected_pages") or []

    if not upload_id:
        return jsonify({"error": "Missing upload id."}), 400

    upload = get_upload(str(upload_id))
    selected_pages = normalize_selected_pages(list(raw_pages), upload.page_count)
    if not selected_pages:
        return jsonify({"error": "Select at least one page."}), 400

    job_id = uuid4().hex
    job = JobRecord(job_id=job_id, upload_id=upload.upload_id, selected_pages=selected_pages)
    with registry_lock:
        jobs[job_id] = job

    append_event(job, "status", status="queued")
    worker = threading.Thread(target=run_transcription_job, args=(job_id,), daemon=True)
    worker.start()

    return jsonify(
        {
            "job_id": job_id,
            "stream_url": f"/api/jobs/{job_id}/stream",
            "status_url": f"/api/jobs/{job_id}",
        }
    )


@app.get("/api/jobs/<job_id>")
def job_status(job_id: str) -> Response:
    return jsonify(get_job_snapshot(get_job(job_id)))


@app.get("/api/jobs/<job_id>/stream")
def stream_job(job_id: str):
    job = get_job(job_id)

    @stream_with_context
    def event_stream():
        index = 0
        while True:
            event = None
            finished = False
            with job.condition:
                while index >= len(job.events) and job.status not in {"completed", "failed"}:
                    job.condition.wait(timeout=15)
                    if index >= len(job.events) and job.status not in {"completed", "failed"}:
                        break
                if index < len(job.events):
                    event = job.events[index]
                    index += 1
                else:
                    finished = job.status in {"completed", "failed"}

            if event is not None:
                yield f"data: {json.dumps(event)}\n\n"
                continue

            if finished:
                break

            yield ": keep-alive\n\n"

    response = Response(event_stream(), mimetype="text/event-stream")
    response.headers["Cache-Control"] = "no-cache"
    response.headers["X-Accel-Buffering"] = "no"
    return response


@app.get("/api/jobs/<job_id>/download/txt")
def download_txt(job_id: str):
    job = get_job(job_id)
    if job.status != "completed":
        abort(409, description="The transcription is not finished yet.")

    upload = get_upload(job.upload_id)
    archive = io.BytesIO()
    with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED) as zip_file:
        for page_number in job.selected_pages:
            page_text = job.page_texts.get(page_number, "")
            zip_file.writestr(f"page_{page_number:03d}.txt", page_text)
    archive.seek(0)

    return send_file(
        archive,
        mimetype="application/zip",
        as_attachment=True,
        download_name=f"{build_safe_stem(upload.original_name)}_pages_txt.zip",
    )


@app.get("/api/jobs/<job_id>/download/docx")
def download_docx(job_id: str):
    job = get_job(job_id)
    if job.status != "completed":
        abort(409, description="The transcription is not finished yet.")

    upload = get_upload(job.upload_id)
    document = Document()
    for index, page_number in enumerate(job.selected_pages):
        lines = job.page_texts.get(page_number, "").split("\n")
        last_paragraph = None
        for line in lines:
            paragraph = document.add_paragraph()
            rtl = _is_rtl(line)
            _apply_paragraph_bidi(paragraph, rtl)
            if line:
                run = paragraph.add_run(line)
                _apply_run_bidi(run, rtl)
            last_paragraph = paragraph
        if index < len(job.selected_pages) - 1 and last_paragraph is not None:
            last_paragraph.add_run().add_break(WD_BREAK.PAGE)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"{build_safe_stem(upload.original_name)}_transcription.docx",
    )


ensure_storage()


if __name__ == "__main__":
    debug_enabled = os.environ.get("FLASK_DEBUG", "0") == "1"
    port = int(os.environ.get("PORT", "5000"))
    app.run(
        host="127.0.0.1",
        port=port,
        debug=debug_enabled,
        use_reloader=debug_enabled,
        threaded=True,
    )